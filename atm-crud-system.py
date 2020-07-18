import sqlite3
import os
import pickle
from docx import Document

def create_account():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        c.execute('CREATE TABLE customer (\
            customerID TEXT NOT NULL PRIMARY KEY,\
            name TEXT NOT NULL,\
            surename TEXT NOT NULL,\
            password TEXT NOT NULL,\
            address TEXT NOT NULL,\
            sex TEXT NOT NULL,\
            age INT NOT NULL,\
            email TEXT NOT NULL,\
            education TEXT NOT NULL,\
            money INT NOT NULL,\
            interest INT NOT NULL\
            )')
        c.execute('CREATE TABLE transactions(\
            customerID TEXT NOT NULL, \
            password TEXT NOT NULL,\
            money INT NOT NULL, \
            deposit INT,\
            withdraw INT\
            )' )
        conn.commit()
        print("Database is connected")
    except Exception as e:
        print("{}".format(e))

def insert_account():
    conn = sqlite3.connect('Account.sqlite')
    try:
        count = True
        c = conn.cursor()
        name = input("Enter your name: ")
        surename = input("Enter your surename: ")
        address = input("Enter your address: ")
        while(count):
            sex = input("Enter your gender[M/F]: ")
            if sex in ('M','F'):
                break
        age = int(input("Enter your age: "))
        email = input("Enter your Email: ")
        education = input("Enter your education: ")
        while(count):
            money = int(input("Enter your money[first service 500 is minimum]: "))
            if(money >= 500):
                break
        interest = int(input("Enter interest rate per year:"))
        cID = input("Enter your ID: ")
        password = input("Enter your password: ")
        data = (cID, name, surename, password, address, sex, age, email, education, money, interest)          
        c.execute('INSERT INTO customer(\
            customerID,\
            name, \
            surename, \
            password, \
            address, \
            sex, \
            age, \
            email, \
            education, \
            money, \
            interest\
            ) VALUES (?,?,?,?,?,?,?,?,?,?,?)', data)
        c.execute('INSERT INTO transactions(customerID, password, money) VALUES (?,?,?)',(cID, password, money))
        print("Account is created")
    except Exception as e:
        print("{}".format(e))

def show_data():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        for row in c.execute('SELECT * FROM customer'):
            print(row)
    except Exception as e:
            print("{}".format(e))

def withdraw():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        c.execute('SELECT password FROM customer')
        passwordtuple = c.fetchall()
        passwordlist = []
        for a in passwordtuple:
            for i in a:
                passwordlist.append(i)
        c.execute('SELECT customerID FROM customer')
        IDtuple = c.fetchall()
        IDlist = []
        for a in IDtuple:
            for i in a:
                IDlist.append(i)
        print("===========Withdraw============")
        cID = input("Enter your ID: ")
        cpassword = input("Enter your password: ")
        while(True):
            if(cpassword in passwordlist and cID in IDlist):
                c.execute('''SELECT money FROM customer WHERE password = ?''',(cpassword))
                tupmon = c.fetchall()
                money = []
                for t in tupmon:
                    for a in t:
                        money.append(a)
                print("Total Money: {} Baht.".format(money[0]))
                withdraw = int(input("How much do you want to withdraw: "))
                while(True):
                    if(withdraw <= money[0] and withdraw != 0):
                        new = money[0] - withdraw
                        c.execute('''UPDATE customer SET money = ? WHERE password = ?''',(new,cpassword))
                        print("Withdraw done")
                        print("Total Money: {} Baht.\nThank you.".format(new))
                        c.execute('''INSERT INTO transactions(
                            customerID,
                            password, 
                            money, 
                            withdraw
                            ) VALUES (?,?,?,?)''',(cID,cpassword,new,withdraw))
                        break
                    elif(withdraw > money[0]):
                        print("You don't have enough money")
                        withdraw = int(input("How much do you want to withdraw: "))
                    elif(withdraw == 0):
                        withdraw = int(input("Invalid value try again: "))
                break
            else:
                cID = input("Your ID/password not match! try again.\nEnter your ID: ")
                cpassword = input("password: ")
    except Exception as e:
        print("{}".format(e))

def deposite():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        c.execute('SELECT password FROM customer')
        passwordtuple = c.fetchall()
        passwordlist = []
        for x in passwordtuple:
            for i in x:
                passwordlist.append(i)
        c.execute('SELECT customerID FROM customer')
        IDtuple = c.fetchall()
        IDlist = []
        for x in IDtuple:
            for i in x:
                IDlist.append(i)
        print("===========Deposite============")
        cID = input("Enter your ID: ")
        cpassword = input("Enter your password: ")
        while(True):
            if(cpassword in passwordlist and cID in IDlist):
                c.execute('SELECT money FROM customer WHERE password = ?',(cpassword))
                tupmon = c.fetchall()
                money = []
                for t in tupmon:
                    for a in t:
                        money.append(a)
                print("Total Money: {} Baht.".format(money[0]))
                dp = int(input('How many do you want to deposite:'))
                while(True):
                    if(dp>0):
                        new = money[0]+dp
                        c.execute('UPDATE customer SET money = ? WHERE password = ?',(new,cpassword))
                        print("Deposite done")
                        print("Total Money: {} Baht.\nThankyou".format(new))
                        c.execute('INSERT INTO transactions(\
                            customerID,\
                            password,\
                            money,\
                            deposite\
                            ) VALUES (?,?,?,?)',(cID,cpassword,new,deposite,))
                        break
                    elif(deposite == 0):
                        deposite = int(input("Invalid value try again: "))
                break
            else:
                cID = input("Your ID/password not match! try again.\nEnter your ID: ")
                cpassword = input("password: ")
    except Exception as e:
        print("{}".format(e))

def daily_transaction():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        a = c.execute('SELECT customerID,\
            sum(deposite),\
            sum(withdraw),\
            money FROM transactions group by customerID')
        s = a.fetchall()
        count = len(s)
        doc = Document()
        t = doc.add_table(rows = count + 1, cols = 5)
        t.style = "Table Grid"
        t.rows[0].cells[0].text = "Customer ID"
        t.rows[0].cells[1].text = "Total Money"
        t.rows[0].cells[2].text = "Total Transaction"
        t.rows[0].cells[3].text = "Total Deposite"
        t.rows[0].cells[4].text = "Total Withdraw"
        for x in range(count):
            if(x < count):
                row = t.rows[x+1]
                row.cells[0].text = "{}".format(s[x][0])
                row.cells[1].text = "{}".format(s[x][3])
                row.cells[2].text = "{}".format(s[x][1]+s[x][2])
                row.cells[3].text = "{}".format(s[x][1])
                row.cells[4].text = "{}".format(s[x][2])
        doc.save("Account.docx")
        c.execute('drop table if exists transactions')
        c.execute('CREATE TABLE transactions(\
            customerID TEXT NOT NULL,\
            password TEXT NOT NULL,\
            money INT NOT NULL,\
            deposite INT,\
            withdraw INT\
            )''')
    except Exception as e:
        print("{}".format(e))

def interest_rate():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        cID = input("Enter your ID: ")
        c.execute('SELECT customerID FROM customer')
        IDtuple = c.fetchall()
        IDlist = []
        for x in IDtuple:
            for i in x:
                    IDlist.append(i)
        while(True):
            if(cID in IDlist):
                c.execute('UPDATE customer SET money = money + (money * (interest/100)) WHERE customerID = ?',(cID))
                a = c.execute('select money from customer where customerID=?',(cID)).fetchall()
                mon = 0
                for i in a:
                    for x in i:
                        mon = x
                print("Total money: {} Baht".format(mon))
                    
                break
            else:
                cID = input("Wrong ! Please enter your ID again: ")
    except Exception as e:
        print("{}".format(e))

def customer_document():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        a = c.execute('SELECT * from customer')
        s = a.fetchall()
        count = len(s)
        doc = Document()
        t = doc.add_table(rows = count + 1, cols = 11)
        t.style = 'Table Grid'
        t.rows[0].cells[0].text = 'Customer ID'
        t.rows[0].cells[1].text = 'Name'
        t.rows[0].cells[2].text = 'Surname'
        t.rows[0].cells[3].text = 'Password'
        t.rows[0].cells[4].text = 'Address'
        t.rows[0].cells[5].text = 'Sex'
        t.rows[0].cells[6].text = 'Age'
        t.rows[0].cells[7].text = 'Email'
        t.rows[0].cells[8].text = 'Education'
        t.rows[0].cells[9].text = 'Money'
        t.rows[0].cells[10].text = 'Interest rate'          
        for x in range(count):
            if(x < count):
                row = t.rows[x + 1]
                row.cells[0].text = "{s[x][0]}".format(s[x][0])
                row.cells[1].text = "{s[x][1]}".format(s[x][1])
                row.cells[2].text = "{s[x][2]}".format(s[x][2])
                row.cells[3].text = "{s[x][3]}".format(s[x][3])
                row.cells[4].text = "{s[x][4]}".format(s[x][4])
                row.cells[5].text = "{s[x][5]}".format(s[x][5])
                row.cells[6].text = "{s[x][6]}".format(s[x][6])
                row.cells[7].text = "{s[x][7]}".format(s[x][7])
                row.cells[8].text = "{s[x][8]}".format(s[x][8])
                row.cells[9].text = "{s[x][9]}".format(s[x][9])
                row.cells[10].text = "{s[x][10]}".format(s[x][10])
        doc.save('Customer.docx')
    except Exception as e:
        print("{}".format(e))

def dump():
    conn = sqlite3.connect('Account.sqlite')
    try:
        c = conn.cursor()
        dump = []
        b = c.execute('select * from customer').fetchall()
        customer_document()
        output = open("Account.bin","wb")
        pickle.dump(b, output)
        print("Dump Finish!")
        output.close()
    except Exception as e:
        print("{}".format(e))
def clear_screen():
    os.system('cls')
    print()

if __name__=='__main__':
    print('>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>>WELCOME<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<<')
    create_account()
    count = True
    while count:
        choice = int(input("\n1) deposite\n2) withdraw\n3) create account\n4) exit\n\nchoice>>> "))
        if choice == 1:
            deposite()
        elif choice == 2:
            withdraw()
        elif choice == 3:
            insert_account()
        elif choice == 4:
            # close the program
            exit()
        else:
            clear_screen()
            print("ERROR: Wrong choice\n")
        #show_data()
        #daily_transaction()
        #interest_rate()
        #customer_document()
        #dump()